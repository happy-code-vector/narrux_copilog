"""KB document ingestion pipeline.
NO pydantic_ai imports. Pure Python.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

import structlog
import tiktoken

from config import get_settings
from retrieval.embeddings import embed_documents
from retrieval.vector_store import (
    delete_chunks_for_document,
    upsert_chunks,
    upsert_document,
)
from tools.schemas import DocumentScope, KBChunk, KBDocument

logger = structlog.get_logger(__name__)

# ─── Data Classes ────────────────────────────────────────────────────────────


@dataclass
class IngestionSource:
    """Definition of a document to ingest."""

    path: Path
    doc_id: str
    doc_version: str
    title: str
    scope: DocumentScope
    owner: str = "NARRUX"
    strategy: str | None = None
    volume: str | None = None
    module_id: str | None = None
    supersedes: str | None = None
    module_markers: dict[str, str] = field(default_factory=dict)


@dataclass
class RawChunk:
    """A raw chunk before embedding."""

    content: str
    metadata: dict


# ─── Text Extraction ─────────────────────────────────────────────────────────


def _is_zip_archive(path: Path) -> bool:
    """Check if a file is a ZIP archive by reading magic bytes.

    Checks for PK\\x03\\x04, NOT file extension.
    """
    try:
        with open(path, "rb") as f:
            magic = f.read(4)
        return magic == b"PK\x03\x04"
    except Exception:
        return False


def extract_text_from_zip(path: Path) -> list[tuple[int, str]]:
    """Extract text from a ZIP archive containing .txt files (one per page).

    Returns list of (page_number, text) sorted by page number.
    """
    pages: list[tuple[int, str]] = []
    with zipfile.ZipFile(path, "r") as zf:
        txt_files = [n for n in zf.namelist() if n.endswith(".txt")]
        for txt_name in txt_files:
            # Parse page number from filename (e.g., "1.txt" → 1)
            try:
                page_num = int(txt_name.replace(".txt", "").split("/")[-1])
            except ValueError:
                page_num = 0

            raw = zf.read(txt_name)
            text = raw.decode("utf-8", errors="replace")
            text = _strip_page_header_footer(text)
            pages.append((page_num, text))

    pages.sort(key=lambda x: x[0])
    return pages


def _strip_page_header_footer(text: str) -> str:
    """Remove NARRUX header/footer patterns from page text."""
    patterns = [
        r"^Strictly Confidential.*$",
        r"^NARRUX.*(?:Strictly Confidential|Handbook|Specification).*$",
        r"^Confidential — internal use only.*$",
        r"^Page \d+ of \d+$",
    ]
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        skip = False
        for pattern in patterns:
            if re.match(pattern, line.strip(), re.IGNORECASE):
                skip = True
                break
        if not skip:
            cleaned.append(line)
    return "\n".join(cleaned)


def extract_text_from_md(path: Path) -> str:
    """Extract text from a markdown file."""
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text_from_txt(path: Path) -> str:
    """Extract text from a plain text file."""
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a .docx file."""
    from docx import Document

    doc = Document(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_csv(path: Path) -> str:
    """Extract text from a CSV file (e.g., input index files)."""
    import csv

    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)


def extract_text_from_json(path: Path) -> str:
    """Extract text from a JSON file (e.g., filter glossary)."""
    import json

    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, list):
        # Array of objects — flatten each to key: value lines
        parts = []
        for item in data:
            if isinstance(item, dict):
                lines = [f"{k}: {v}" for k, v in item.items() if v]
                parts.append("\n".join(lines))
            else:
                parts.append(str(item))
        return "\n\n".join(parts)
    elif isinstance(data, dict):
        # Object — flatten to key: value lines
        lines = []
        for k, v in data.items():
            if isinstance(v, (list, dict)):
                lines.append(f"{k}: {json.dumps(v)}")
            else:
                lines.append(f"{k}: {v}")
        return "\n".join(lines)
    return str(data)


# ─── Chunking ────────────────────────────────────────────────────────────────


def _count_tokens(text: str) -> int:
    """Count tokens using tiktoken cl100k_base."""
    enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))


def make_chunk_id(doc_id: str, content: str, index: int) -> str:
    """Generate a deterministic chunk ID."""
    hash_input = f"{doc_id}::{index}::{content[:200]}"
    hash_hex = hashlib.sha256(hash_input.encode()).hexdigest()[:16]
    return f"{doc_id}::{hash_hex}"


def chunk_by_module_boundaries(
    pages: list[tuple[int, str]], source: IngestionSource
) -> list[RawChunk]:
    """Chunk by module boundaries (e.g., D1 · CVD Filter [C]).

    Falls back to chunk_sliding_window if no module headers found.
    """
    settings = get_settings()
    full_text = "\n\n".join(text for _, text in pages)

    # Module header pattern: A1 · Module Name [A]
    pattern = re.compile(
        r"^([A-L]\d{1,2})\s*[·•]\s*(.+?)(?:\s*\[([ABC])\])?$",
        re.MULTILINE,
    )

    matches = list(pattern.finditer(full_text))
    if not matches:
        logger.info("no_module_headers_found", doc_id=source.doc_id, fallback="sliding_window")
        return chunk_sliding_window(full_text, source)

    logger.info("module_headers_found", doc_id=source.doc_id, count=len(matches))

    chunks: list[RawChunk] = []

    for i, match in enumerate(matches):
        module_id = match.group(1)
        module_name = match.group(2).strip()
        param_class = match.group(3)

        # Extract section text (from this match to next match or end)
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        section_text = full_text[start:end].strip()

        meta: dict = {
            "source_type": "handbook",
            "strategy": source.strategy,
            "volume": source.volume,
            "module_id": module_id,
            "module_name": module_name,
            "doc_id": source.doc_id,
            "doc_version": source.doc_version,
        }
        if param_class:
            meta["param_class"] = param_class
            meta["regime_coupled"] = param_class == "C"

        section_tokens = _count_tokens(section_text)

        if section_tokens > settings.chunk_size_tokens * 2:
            # Subdivide large sections with sliding window
            sub_chunks = chunk_sliding_window(section_text, source, extra_meta=meta)
            chunks.extend(sub_chunks)
        else:
            chunks.append(RawChunk(content=section_text, metadata=meta))

    logger.info("module_chunking_complete", doc_id=source.doc_id, chunks=len(chunks))
    return chunks


def chunk_sliding_window(
    text: str,
    source: IngestionSource,
    extra_meta: dict | None = None,
) -> list[RawChunk]:
    """Sliding window chunking with tiktoken."""
    settings = get_settings()
    enc = tiktoken.get_encoding("cl100k_base")
    tokens = enc.encode(text)

    chunk_size = settings.chunk_size_tokens
    overlap = settings.chunk_overlap_tokens

    chunks: list[RawChunk] = []
    start = 0
    index = 0

    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = enc.decode(chunk_tokens)

        meta: dict = {"doc_id": source.doc_id, "doc_version": source.doc_version, "chunk_index": index}
        if extra_meta:
            meta.update(extra_meta)

        chunks.append(RawChunk(content=chunk_text, metadata=meta))

        start = end - overlap
        index += 1
        if end >= len(tokens):
            break

    return chunks


def chunk_pages_with_context(
    pages: list[tuple[int, str]], source: IngestionSource
) -> list[RawChunk]:
    """One chunk per page if fits in chunk_size. Subdivide if too long."""
    settings = get_settings()
    chunks: list[RawChunk] = []

    for page_num, text in pages:
        if not text.strip():
            continue

        token_count = _count_tokens(text)

        if token_count <= settings.chunk_size_tokens:
            meta = {
                "doc_id": source.doc_id,
                "doc_version": source.doc_version,
                "page_number": page_num,
            }
            chunks.append(RawChunk(content=text, metadata=meta))
        else:
            # Subdivide with sliding window
            sub_chunks = chunk_sliding_window(
                text, source, extra_meta={"page_number": page_num}
            )
            chunks.extend(sub_chunks)

    return chunks


# ─── Ingestion Pipeline ──────────────────────────────────────────────────────


async def ingest_document(source: IngestionSource, force_reingest: bool = False) -> int:
    """Ingest a single document into the knowledge base.

    Steps:
    1. Detect file type with _is_zip_archive() first.
    2. Choose chunking strategy.
    3. Register document via upsert_document.
    4. Delete old chunks (always, to handle version upgrades).
    5. Embed with embed_documents.
    6. Build KBChunk objects with make_chunk_id.
    7. Upsert chunks. Return count.
    """
    path = source.path
    if not path.exists():
        logger.warning("file_not_found", path=str(path))
        return 0

    logger.info("ingesting", doc_id=source.doc_id, path=str(path))

    # Step 1: Extract text
    if _is_zip_archive(path):
        pages = extract_text_from_zip(path)
        text = "\n\n".join(t for _, t in pages)
        use_pages = True
    else:
        suffix = path.suffix.lower()
        if suffix == ".md":
            text = extract_text_from_md(path)
        elif suffix == ".txt":
            text = extract_text_from_txt(path)
        elif suffix == ".docx":
            text = extract_text_from_docx(path)
        elif suffix == ".csv":
            text = extract_text_from_csv(path)
        elif suffix == ".json":
            text = extract_text_from_json(path)
        else:
            text = path.read_text(encoding="utf-8", errors="replace")
        pages = [(1, text)]
        use_pages = True

    if not text.strip():
        logger.warning("empty_document", doc_id=source.doc_id)
        return 0

    # Step 2: Choose chunking strategy
    if source.strategy and source.volume:
        # Handbook → chunk by module boundaries
        raw_chunks = chunk_by_module_boundaries(pages, source)
    elif source.scope in (DocumentScope.governance, DocumentScope.process):
        # Spec/process → chunk pages with context
        raw_chunks = chunk_pages_with_context(pages, source)
    else:
        # Default → sliding window
        raw_chunks = chunk_sliding_window(text, source)

    if not raw_chunks:
        logger.warning("no_chunks_generated", doc_id=source.doc_id)
        return 0

    # Step 3: Register document
    doc = KBDocument(
        doc_id=source.doc_id,
        doc_version=source.doc_version,
        title=source.title,
        scope=source.scope,
        strategy=source.strategy,
        volume=source.volume,
        module_id=source.module_id,
        owner=source.owner,
        supersedes=source.supersedes,
    )
    await upsert_document(doc)

    # Step 4: Delete old chunks
    await delete_chunks_for_document(source.doc_id)

    # Step 5: Embed
    texts = [c.content for c in raw_chunks]
    embeddings = await embed_documents(texts)

    logger.debug("embedding_result", raw_chunks=len(raw_chunks), embeddings=len(embeddings))

    # Step 6: Build KBChunk objects
    kb_chunks: list[KBChunk] = []
    for i, (raw, embedding) in enumerate(zip(raw_chunks, embeddings)):
        chunk_id = make_chunk_id(source.doc_id, raw.content, i)
        kb_chunks.append(
            KBChunk(
                chunk_id=chunk_id,
                doc_id=source.doc_id,
                doc_version=source.doc_version,
                content=raw.content,
                token_count=_count_tokens(raw.content),
                embedding=embedding,
                metadata=raw.metadata,
            )
        )

    logger.debug("kb_chunks_built", count=len(kb_chunks))

    # Step 7: Upsert
    count = await upsert_chunks(kb_chunks)
    logger.info("ingestion_complete", doc_id=source.doc_id, chunks=count)
    return count


async def ingest_all_project_documents(kb_dir: Path) -> dict[str, int]:
    """Ingest all project documents defined in the source registry.

    Args:
        kb_dir: Path to kb_content/ directory.

    Returns:
        Dict of {doc_id: chunk_count} for each document.
    """
    sources = _build_source_registry(kb_dir)
    results: dict[str, int] = {}

    for source in sources:
        if not source.path.exists():
            logger.warning("source_not_found", doc_id=source.doc_id, path=str(source.path))
            results[source.doc_id] = -1
            continue

        try:
            count = await ingest_document(source)
            results[source.doc_id] = count
        except Exception as e:
            logger.error("ingestion_error", doc_id=source.doc_id, error=str(e))
            results[source.doc_id] = -1

    return results


def _build_source_registry(kb_dir: Path) -> list[IngestionSource]:
    """Build the registry of all documents to ingest.

    kb_content/ layout:
      strategies/     → strategy handbooks (per-strategy subdirs)
      filters/        → filter glossary + JSON lookup
      parameters/     → input index CSVs, bounds (future)
      governance/     → specs, frameworks, rules
      playbook/       → edge cases, function specs
      templates/      → report templates, reference appendices
    """
    sources: list[IngestionSource] = []

    # ─── Strategy handbooks ───────────────────────────────────────────────────
    strategies_dir = kb_dir / "strategies"

    # Alpha handbooks (5 volumes)
    alpha_dir = strategies_dir / "Alpha"
    alpha_volumes = [
        ("alpha_handbook_v15_9_1_vol_ab", "Vol_A-B", "AB"),
        ("alpha_handbook_v15_9_1_vol_c", "Vol_C", "C"),
        ("alpha_handbook_v15_9_1_vol_d", "Vol_D", "D"),
        ("alpha_handbook_v15_9_1_vol_ef", "Vol_E-F", "EF"),
        ("alpha_handbook_v15_9_1_vol_hl", "Vol_H-L", "HL"),
    ]
    for doc_id, vol_tag, volume in alpha_volumes:
        filename = f"NARRUX_Alpha_v15.9.1_Handbook_{vol_tag}.md"
        sources.append(IngestionSource(
            path=alpha_dir / filename, doc_id=doc_id, doc_version="15.9.1",
            title=f"Alpha Handbook {vol_tag}", scope=DocumentScope.strategy,
            strategy="alpha", volume=volume,
        ))

    # Sentinel handbook
    sources.append(IngestionSource(
        path=strategies_dir / "Sentinel" / "NARRUX_Sentinel_v1.9_Handbook.md",
        doc_id="sentinel_handbook_v1_9", doc_version="1.9",
        title="NARRUX Sentinel v1.9 Handbook", scope=DocumentScope.strategy,
        strategy="sentinel",
    ))

    # Master handbook
    sources.append(IngestionSource(
        path=strategies_dir / "MAster" / "NARRUX_Master_v14.3_Handbook.md",
        doc_id="master_handbook_v14_3", doc_version="14.3",
        title="NARRUX Master v14.3 Handbook", scope=DocumentScope.strategy,
        strategy="master",
    ))

    # NRX handbook
    sources.append(IngestionSource(
        path=strategies_dir / "NRX" / "NARRUX_NRX_MTrv1_Handbook.md",
        doc_id="nrx_mtr_v1_handbook", doc_version="1.0",
        title="NARRUX NRX MTrv1 Handbook", scope=DocumentScope.strategy,
        strategy="nrx",
    ))

    # Strategy comparison matrix
    sources.append(IngestionSource(
        path=strategies_dir / "NARRUX_Strategy_Comparison_Matrix_v1.0.md",
        doc_id="strategy_comparison_matrix_v1_0", doc_version="1.0",
        title="NARRUX Strategy Comparison Matrix v1.0", scope=DocumentScope.strategy,
    ))

    # ─── Input index CSVs ─────────────────────────────────────────────────────
    params_dir = kb_dir / "parameters"
    input_indices = [
        ("alpha_v15_9_1_input_index", "alpha_v15_9_1_input_index.csv", "15.9.1", "alpha"),
        ("sentinel_v1_9_input_index", "sentinel_v1_9_input_index.csv", "1.9", "sentinel"),
        ("master_v14_3_input_index", "master_v14_3_input_index.csv", "14.3", "master"),
        ("nrx_mtr_v1_input_index", "nrx_mtr_v1_input_index.csv", "1.0", "nrx"),
    ]
    for doc_id, filename, ver, strategy in input_indices:
        sources.append(IngestionSource(
            path=params_dir / filename, doc_id=doc_id, doc_version=ver,
            title=f"{strategy.title()} Input Index", scope=DocumentScope.strategy,
            strategy=strategy,
        ))

    # ─── Filter glossary ──────────────────────────────────────────────────────
    filters_dir = kb_dir / "filters"
    sources.append(IngestionSource(
        path=filters_dir / "NARRUX_Filter_Glossary_and_Param_Classes_v1.1.md",
        doc_id="filter_glossary_and_param_classes_v1_1", doc_version="1.1",
        title="NARRUX Filter Glossary and Param Classes v1.1",
        scope=DocumentScope.filter_glossary,
    ))
    sources.append(IngestionSource(
        path=filters_dir / "narrux_filter_glossary.json",
        doc_id="filter_glossary_json", doc_version="1.0",
        title="NARRUX Filter Glossary (JSON)", scope=DocumentScope.filter_glossary,
    ))

    # ─── Governance ───────────────────────────────────────────────────────────
    gov_dir = kb_dir / "governance"
    governance_docs = [
        ("tsi_v2_spec", "2.0", "NARRUX TSI v2.0 CA Engineering Spec",
         "NARRUX_TSI_v2.0_CA_Engineering_Spec.md", DocumentScope.governance),
        ("leverage_framework_v1_0", "1.0", "NARRUX Leverage Framework v1.0",
         "NARRUX_Leverage_Framework_v1.0.md", DocumentScope.governance),
        ("metric_definitions_v1_0", "1.0", "NARRUX Metric Definitions v1.0",
         "NARRUX_Metric_Definitions_v1.0.md", DocumentScope.governance),
        ("backtest_analysis_approach_v1_0", "1.0", "NARRUX Backtest Analysis Approach v1.0",
         "NARRUX_Backtest_Analysis_Approach_v1.0.md", DocumentScope.process),
        ("kb_routing_and_guardrails_v1_0", "1.0", "NARRUX KB Routing and Guardrails v1.0",
         "NARRUX_CoPilot_KB_Routing_and_Guardrails_v1.0.md", DocumentScope.governance),
    ]
    for doc_id, ver, title, filename, scope in governance_docs:
        sources.append(IngestionSource(
            path=gov_dir / filename, doc_id=doc_id, doc_version=ver,
            title=title, scope=scope,
        ))

    # ─── Playbook ─────────────────────────────────────────────────────────────
    playbook_dir = kb_dir / "playbook"
    playbook_docs = [
        ("edgecase_playbook_v1_0", "1.0", "NARRUX EdgeCase Playbook v1.0",
         "NARRUX_EdgeCase_Playbook_v1.0.md", DocumentScope.playbook),
        ("drift_monitor_f05_spec_v1_0", "1.0", "NARRUX Drift Monitor F05 Spec v1.0",
         "NARRUX_Drift_Monitor_F05_Spec_v1.0.md", DocumentScope.governance),
    ]
    for doc_id, ver, title, filename, scope in playbook_docs:
        sources.append(IngestionSource(
            path=playbook_dir / filename, doc_id=doc_id, doc_version=ver,
            title=title, scope=scope,
        ))

    # ─── Templates ────────────────────────────────────────────────────────────
    templates_dir = kb_dir / "templates"
    templates_docs = [
        ("copilot_report_template_v1_0", "1.0", "NARRUX CoPilot Report Template v1.0",
         "NARRUX_CoPilot_Report_Template_v1.0.md", DocumentScope.report_template),
        ("input_index_appendix", "1.0", "NARRUX Input Index Appendix",
         "NARRUX_Input_Index_Appendix.md", DocumentScope.strategy),
    ]
    for doc_id, ver, title, filename, scope in templates_docs:
        sources.append(IngestionSource(
            path=templates_dir / filename, doc_id=doc_id, doc_version=ver,
            title=title, scope=scope,
        ))

    # ─── Functional spec (at project root) ────────────────────────────────────
    spec_path = kb_dir.parent.parent / "docs" / "aiTrate_AI_Agent_Functional_Spec_v1.0.pdf"
    if spec_path.exists():
        sources.append(IngestionSource(
            path=spec_path,
            doc_id="aitrate_agent_spec_v1_0", doc_version="1.0",
            title="aiTrate AI Agent Functional Spec v1.0", scope=DocumentScope.governance,
        ))

    return sources
